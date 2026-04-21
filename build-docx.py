"""Convert the session transcript markdown to a Word .docx file."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = Path(r"C:\Paul\navigator-crm-app\Roadrunner-CRM-Session-Transcript.md")
OUT = Path(r"C:\Paul\navigator-crm-app\Roadrunner-CRM-Session-Transcript.docx")

doc = Document()

# Base style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

lines = SRC.read_text(encoding='utf-8').splitlines()
i = 0
in_code = False
code_buf = []
in_table = False
table_rows = []

def flush_code():
    global code_buf
    if not code_buf:
        return
    p = doc.add_paragraph()
    run = p.add_run('\n'.join(code_buf))
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.3)
    code_buf = []

def flush_table():
    global table_rows
    if not table_rows:
        return
    # table_rows is list of lists
    rows = table_rows
    # filter separator rows (all dashes)
    rows = [r for r in rows if not all(set(c.strip()) <= set('-: ') for c in r)]
    if not rows:
        table_rows = []
        return
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = 'Light Grid Accent 1'
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            val = row[ci] if ci < len(row) else ''
            cell = tbl.rows[ri].cells[ci]
            cell.text = val.strip()
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    if ri == 0:
                        run.bold = True
    doc.add_paragraph()
    table_rows = []

def add_inline(paragraph, text):
    """Render simple **bold**, *italic*, `code` within a paragraph."""
    # Split by code spans first
    parts = re.split(r'(`[^`]+`)', text)
    for part in parts:
        if part.startswith('`') and part.endswith('`') and len(part) >= 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        else:
            # bold **
            sub = re.split(r'(\*\*[^*]+\*\*)', part)
            for s in sub:
                if s.startswith('**') and s.endswith('**'):
                    run = paragraph.add_run(s[2:-2])
                    run.bold = True
                else:
                    # italic *
                    subi = re.split(r'(\*[^*]+\*)', s)
                    for si in subi:
                        if si.startswith('*') and si.endswith('*') and len(si) >= 2:
                            run = paragraph.add_run(si[1:-1])
                            run.italic = True
                        else:
                            paragraph.add_run(si)

while i < len(lines):
    line = lines[i]

    # code fence
    if line.strip().startswith('```'):
        if in_code:
            flush_code()
            in_code = False
        else:
            flush_table()
            in_code = True
        i += 1
        continue

    if in_code:
        code_buf.append(line)
        i += 1
        continue

    # table row
    if line.strip().startswith('|') and line.strip().endswith('|'):
        parts = [c for c in line.strip().strip('|').split('|')]
        table_rows.append(parts)
        i += 1
        continue
    else:
        if in_table or table_rows:
            flush_table()

    # horizontal rule
    if line.strip() == '---':
        doc.add_paragraph().add_run().add_break()
        i += 1
        continue

    # headings
    m = re.match(r'^(#{1,6})\s+(.*)', line)
    if m:
        level = len(m.group(1))
        text = m.group(2).strip()
        h = doc.add_heading(level=min(level, 4))
        # Clean inline markdown from heading
        clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        clean = re.sub(r'`([^`]+)`', r'\1', clean)
        h.add_run(clean)
        i += 1
        continue

    # bullet
    if re.match(r'^\s*[-*]\s+', line):
        text = re.sub(r'^\s*[-*]\s+', '', line)
        p = doc.add_paragraph(style='List Bullet')
        add_inline(p, text)
        i += 1
        continue

    # numbered list
    if re.match(r'^\s*\d+\.\s+', line):
        text = re.sub(r'^\s*\d+\.\s+', '', line)
        p = doc.add_paragraph(style='List Number')
        add_inline(p, text)
        i += 1
        continue

    # blank
    if line.strip() == '':
        doc.add_paragraph()
        i += 1
        continue

    # regular paragraph
    p = doc.add_paragraph()
    add_inline(p, line)
    i += 1

# flush any tail
flush_code()
flush_table()

doc.save(OUT)
print(f"Wrote: {OUT}")
print(f"Size: {OUT.stat().st_size:,} bytes")
